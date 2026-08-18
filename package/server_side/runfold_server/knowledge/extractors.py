from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from runfold_server.errors import ApiError

TEXT = "text/plain"
MARKDOWN = "text/markdown"
PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def media_type_for(path: Path, original_filename: str) -> str:
    extension = Path(original_filename).suffix.lower()
    prefix = _prefix(path)
    if extension == ".txt":
        if prefix.startswith((b"%PDF-", b"PK\x03\x04")):
            raise _invalid_file()
        return TEXT
    if extension == ".md":
        if prefix.startswith((b"%PDF-", b"PK\x03\x04")):
            raise _invalid_file()
        return MARKDOWN
    if extension == ".pdf":
        if not prefix.startswith(b"%PDF-"):
            raise _invalid_file()
        return PDF
    if extension == ".docx":
        if not prefix.startswith(b"PK\x03\x04"):
            raise _invalid_file()
        return DOCX
    raise ApiError(415, "unsupported_document_type", "Document type is not supported")


def stored_media_type(path: Path, original_filename: str) -> str:
    prefix = _prefix(path)
    if prefix.startswith(b"%PDF-"):
        return PDF
    if prefix.startswith(b"PK\x03\x04"):
        return DOCX
    return MARKDOWN if Path(original_filename).suffix.lower() == ".md" else TEXT


def extract_text(
    path: Path,
    *,
    media_type: str,
    max_characters: int,
    pdf_max_pages: int,
    docx_max_uncompressed_bytes: int,
) -> str:
    try:
        if media_type in {TEXT, MARKDOWN}:
            text = path.read_text(encoding="utf-8")
        elif media_type == PDF:
            text = _extract_pdf(path, pdf_max_pages)
        elif media_type == DOCX:
            text = _extract_docx(path, docx_max_uncompressed_bytes)
        else:
            raise ApiError(415, "unsupported_document_type", "Document type is not supported")
    except ApiError:
        raise
    except Exception as error:
        raise _invalid_file() from error

    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise ApiError(422, "empty_document", "Document contains no extractable text")
    if len(normalized) > max_characters:
        raise ApiError(413, "extracted_text_too_large", "Extracted text exceeds the limit")
    return normalized


def _extract_pdf(path: Path, max_pages: int) -> str:
    reader = PdfReader(path, strict=True)
    if reader.is_encrypted:
        raise _invalid_file()
    if len(reader.pages) > max_pages:
        raise ApiError(413, "pdf_page_limit_exceeded", "PDF page count exceeds the limit")
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(path: Path, max_uncompressed_bytes: int) -> str:
    with zipfile.ZipFile(path) as archive:
        total = sum(item.file_size for item in archive.infolist())
        names = frozenset(item.filename for item in archive.infolist())
        if total > max_uncompressed_bytes:
            raise ApiError(
                413,
                "docx_uncompressed_limit_exceeded",
                "DOCX expanded size exceeds the limit",
            )
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise _invalid_file()

    document = DocxDocument(path)
    values = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            value = "\t".join(cell.text for cell in row.cells).strip()
            if value:
                values.append(value)
    return "\n".join(values)


def _invalid_file() -> ApiError:
    return ApiError(422, "invalid_document", "Document content is invalid")


def _prefix(path: Path) -> bytes:
    with path.open("rb") as source:
        return source.read(8)
