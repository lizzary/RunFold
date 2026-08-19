from __future__ import annotations

import asyncio
from pathlib import Path

import httpx
import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from runfold_server.errors import ApiError
from runfold_server.knowledge.chunker import chunk_text
from runfold_server.knowledge.extractors import DOCX, PDF, extract_text, media_type_for
from runfold_server.knowledge.lance_index import LanceIndex
from runfold_server.llm.openai_embeddings import OpenAIEmbeddingsClient


def test_chunker_prefers_boundaries_respects_size_and_uses_contiguous_ordinals() -> None:
    chunks = chunk_text(
        "Alpha words here.\n\nBeta words are in another paragraph.\n\nGamma finishes.",
        size=32,
        overlap=5,
    )

    assert chunks
    assert [chunk.ordinal for chunk in chunks] == list(range(len(chunks)))
    assert all(0 < len(chunk.text) <= 32 for chunk in chunks)
    assert len({chunk.chunk_id for chunk in chunks}) == len(chunks)


def test_docx_paragraphs_and_tables_are_extracted_with_expanded_size_limit(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Paragraph value")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Left"
    table.cell(0, 1).text = "Right"
    document.save(path)

    assert media_type_for(path, path.name) == DOCX
    text = extract_text(
        path,
        media_type=DOCX,
        max_characters=1_000,
        pdf_max_pages=10,
        docx_max_uncompressed_bytes=1_000_000,
    )
    assert "Paragraph value" in text
    assert "Left\tRight" in text
    with pytest.raises(ApiError, match="expanded size") as error:
        extract_text(
            path,
            media_type=DOCX,
            max_characters=1_000,
            pdf_max_pages=10,
            docx_max_uncompressed_bytes=1,
        )
    assert error.value.status_code == 413


def test_pdf_signature_page_limit_and_empty_text_are_safe(tmp_path: Path) -> None:
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as target:
        writer.write(target)

    assert media_type_for(path, path.name) == PDF
    with pytest.raises(ApiError) as page_error:
        extract_text(
            path,
            media_type=PDF,
            max_characters=1_000,
            pdf_max_pages=1,
            docx_max_uncompressed_bytes=1_000,
        )
    assert page_error.value.code == "pdf_page_limit_exceeded"
    with pytest.raises(ApiError) as empty_error:
        extract_text(
            path,
            media_type=PDF,
            max_characters=1_000,
            pdf_max_pages=2,
            docx_max_uncompressed_bytes=1_000,
        )
    assert empty_error.value.code == "empty_document"


def test_embedding_client_validates_order_dimensions_finiteness_and_usage() -> None:
    async def run() -> None:
        responses = [
            httpx.Response(503),
            httpx.Response(
                200,
                json={
                    "data": [{"index": 0, "embedding": [0.1, 0.2]}],
                    "usage": {"total_tokens": 4},
                },
            ),
        ]

        async def handler(request: httpx.Request) -> httpx.Response:
            assert request.headers["Authorization"] == "Bearer secret"
            return responses.pop(0)

        async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as http:
            client = OpenAIEmbeddingsClient(
                http_client=http,
                base_url="https://example.test/v1",
                api_key="secret",
                model="embedding",
                dimensions=2,
                max_retries=1,
            )
            result = await client.embed(("safe input",))
            assert result.vectors == ((0.1, 0.2),)
            assert result.total_tokens == 4

        async def invalid_handler(_: httpx.Request) -> httpx.Response:
            return httpx.Response(
                200,
                content=(
                    b'{"data":[{"index":0,"embedding":[0.1,NaN]}],'
                    b'"usage":{"total_tokens":1}}'
                ),
                headers={"Content-Type": "application/json"},
            )

        async with httpx.AsyncClient(transport=httpx.MockTransport(invalid_handler)) as http:
            client = OpenAIEmbeddingsClient(
                http_client=http,
                base_url="https://example.test/v1",
                api_key="",
                model="embedding",
                dimensions=2,
                max_retries=0,
            )
            with pytest.raises(ApiError) as error:
                await client.embed(("safe input",))
            assert error.value.code == "invalid_embedding_response"

    asyncio.run(run())


def test_lance_search_prefilters_each_batch_and_merges_global_nearest_hits(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    first_id = "00000000-0000-4000-8000-000000000101"
    second_id = "00000000-0000-4000-8000-000000000102"
    first_chunks = chunk_text("first", size=20, overlap=0)
    second_chunks = chunk_text("second", size=20, overlap=0)
    index = LanceIndex(tmp_path / "lance-search", 2)
    index.recreate()
    asyncio.run(
        index.replace_document(first_id, "1" * 64, first_chunks, ((1.0, 0.0),))
    )
    asyncio.run(
        index.replace_document(second_id, "2" * 64, second_chunks, ((0.0, 1.0),))
    )
    monkeypatch.setattr("runfold_server.knowledge.lance_index._FILTER_BATCH_SIZE", 1)

    hits = index.search(
        (0.0, 1.0), document_ids=(first_id, second_id), top_k=1
    )

    assert len(hits) == 1
    assert hits[0].document_id == second_id
    assert hits[0].text == "second"
